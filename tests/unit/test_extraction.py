"""Unit tests for BE-04 Document Extraction.
Named by requirement IDs per BE-16-R4.
"""

from pathlib import Path
import pytest
import fitz
import docx

from app.core.config import Settings
from app.core.errors import (
    DocumentCorruptError,
    DocumentEncryptedError,
    DocumentTooLongError,
    NoTextExtractedError,
)
from app.ingestion.extract_pdf import extract_pdf
from app.ingestion.extract_docx import extract_docx


def get_test_settings():
    return Settings(
        DATABASE_URL="postgresql://postgres:test@localhost:5432/test",
        SUPABASE_URL="https://test.supabase.co",
        SUPABASE_SERVICE_KEY="test_key",
        GEMINI_API_KEY="test_key",
        MIN_CHARS_PER_PAGE=80,
        MAX_PAGES=500,
    )


def test_be_04_r6_scanned_pdf_below_80_chars_rejected():
    """BE-04-R6: Scanned / image-only PDF with <80 chars/page is rejected."""
    settings = get_test_settings()
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Short text", fontsize=10)  # Only 10 chars
    pdf_bytes = doc.tobytes()
    doc.close()

    with pytest.raises(NoTextExtractedError):
        extract_pdf(pdf_bytes, "scanned.pdf", settings)


def test_be_04_r8_running_headers_suppressed():
    """BE-04-R8: Repeated running headers in top 8% on >50% pages are suppressed."""
    settings = get_test_settings()
    doc = fitz.open()
    
    # 4 pages with identical top header
    for p in range(4):
        page = doc.new_page()
        page.insert_text((50, 30), "ACME CONFIDENTIAL HEADER", fontsize=8)  # top 8%
        page.insert_text((50, 150), f"Body paragraph for page {p+1} with sufficient length to pass density check.", fontsize=11)
        page.insert_text((50, 200), "Additional operational details on enterprise compliance and security standards.", fontsize=11)

    pdf_bytes = doc.tobytes()
    doc.close()

    extracted = extract_pdf(pdf_bytes, "doc_with_headers.pdf", settings)
    assert "ACME CONFIDENTIAL HEADER" not in extracted.text
    assert "Body paragraph for page 1" in extracted.text


def test_be_04_r14_pdf_over_500_pages_rejected():
    """BE-04-R14: PDF exceeding MAX_PAGES (500) fails with DocumentTooLongError."""
    settings = get_test_settings()
    settings.MAX_PAGES = 5  # Small limit for test

    doc = fitz.open()
    for _ in range(6):
        page = doc.new_page()
        page.insert_text((50, 50), "Valid text content for page testing.", fontsize=10)
    pdf_bytes = doc.tobytes()
    doc.close()

    with pytest.raises(DocumentTooLongError):
        extract_pdf(pdf_bytes, "long.pdf", settings)


def test_be_04_r17_docx_tables_converted_to_markdown():
    """BE-04-R17: DOCX tables are converted into Markdown table format."""
    settings = get_test_settings()
    doc = docx.Document()
    doc.add_heading("SLA Matrix", level=1)
    
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Tier"
    table.rows[0].cells[1].text = "Uptime"
    table.rows[1].cells[0].text = "Gold"
    table.rows[1].cells[1].text = "99.95%"

    out_io = Path("test_table.docx")
    doc.save(str(out_io))
    docx_bytes = out_io.read_bytes()
    out_io.unlink(missing_ok=True)

    extracted = extract_docx(docx_bytes, "table.docx", settings)
    assert "| Tier | Uptime |" in extracted.text
    assert "| Gold | 99.95% |" in extracted.text
