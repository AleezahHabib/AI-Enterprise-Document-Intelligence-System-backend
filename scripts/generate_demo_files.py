"""Generate rich, realistic PDF and DOCX files for the demo corpus and test fixtures."""

import os
from pathlib import Path
import fitz  # PyMuPDF
import docx

def create_demo_corpus():
    out_dir = Path(__file__).parent.parent / "data" / "demo_corpus"
    out_dir.mkdir(parents=True, exist_ok=True)

    # 1. Data Protection Policy 2026 (PDF)
    pdf_path = out_dir / "data-protection-policy-2026.pdf"
    doc = fitz.open()
    
    # Page 1: Overview & Scope
    page1 = doc.new_page()
    page1.insert_text((50, 50), "ACME Global Enterprise — Policy Document", fontsize=8, fontname="helv")
    page1.insert_text((50, 80), "Data Protection and Customer Record Retention Policy 2026", fontsize=16, fontname="hebo")
    page1.insert_text((50, 110), "1. Executive Scope and Purpose", fontsize=13, fontname="hebo")
    p1_text = (
        "This policy governs the retention, archival, and cryptographic protection of all customer "
        "records and telemetry collected across ACME enterprise systems. Compliance with this policy "
        "is mandatory for all operating divisions and external service contractors."
    )
    page1.insert_textbox(fitz.Rect(50, 125, 550, 180), p1_text, fontsize=10, fontname="helv")

    page1.insert_text((50, 200), "2. Data Classification Standards", fontsize=13, fontname="hebo")
    p2_text = (
        "All data processed within the organization is categorized into three strict security tiers: "
        "Public Information, Internal Operational Telemetry, and Highly Confidential Customer PII. "
        "Customer PII encompasses government identifiers, banking coordinates, biometric templates, "
        "and primary billing addresses."
    )
    page1.insert_textbox(fitz.Rect(50, 215, 550, 280), p2_text, fontsize=10, fontname="helv")
    page1.insert_text((50, 780), "Confidential — Internal Use Only | Page 1", fontsize=8, fontname="helv")

    # Page 2: Record Retention & Breach Notification
    page2 = doc.new_page()
    page2.insert_text((50, 50), "ACME Global Enterprise — Policy Document", fontsize=8, fontname="helv")
    page2.insert_text((50, 80), "3. Mandatory Retention Schedules", fontsize=13, fontname="hebo")
    page2.insert_text((50, 105), "3.1 Customer Record Retention Period", fontsize=11, fontname="hebo")
    p3_text = (
        "Customer account records and signed transaction logs must be securely retained for exactly "
        "seven (7) years following formal account closure or contract termination. Upon expiry of the "
        "seven-year retention window, all associated primary storage keys and secondary backups must be "
        "cryptographically shredded in accordance with NIST SP 800-88 standards."
    )
    page2.insert_textbox(fitz.Rect(50, 120, 550, 200), p3_text, fontsize=10, fontname="helv")

    page2.insert_text((50, 220), "3.2 Incident Notification SLA", fontsize=11, fontname="hebo")
    p4_text = (
        "In the event of a verified or suspected security breach involving unauthorized access to "
        "customer records, the Security Operations Center (SOC) must notify affected clients and relevant "
        "supervisory authorities within seventy-two (72) hours of initial incident confirmation."
    )
    page2.insert_textbox(fitz.Rect(50, 235, 550, 310), p4_text, fontsize=10, fontname="helv")
    page2.insert_text((50, 780), "Confidential — Internal Use Only | Page 2", fontsize=8, fontname="helv")

    doc.save(str(pdf_path))
    doc.close()
    print(f"Created {pdf_path}")

    # 2. Master Services Agreement (DOCX)
    docx_path = out_dir / "cloud_services_vendor_agreement.docx"
    doc_docx = docx.Document()
    doc_docx.add_heading("Master Cloud Services Agreement", level=1)
    doc_docx.add_paragraph(
        "This Master Cloud Services Agreement ('Agreement') is entered into between CloudProvider Corp ('Vendor') "
        "and ACME Global Enterprise ('Customer') regarding dedicated compute infrastructure and API hosting."
    )

    doc_docx.add_heading("1. Service Level Commitments and Uptime", level=2)
    doc_docx.add_paragraph(
        "Vendor guarantees a monthly uptime service level of 99.95% across all production API gateways. "
        "If monthly availability drops below 99.95% but remains above 99.0%, Customer is entitled to a 15% service credit. "
        "If monthly availability drops below 99.0%, Customer is entitled to a 30% service credit."
    )

    doc_docx.add_heading("2. Dispute Resolution and Arbitration", level=2)
    doc_docx.add_paragraph(
        "Any dispute, controversy, or claim arising out of or relating to this contract, or the breach, "
        "termination, or invalidity thereof, shall be settled by binding arbitration in New York City "
        "administered by the American Arbitration Association in accordance with its Commercial Arbitration Rules."
    )

    doc_docx.add_heading("3. Liability Caps", level=2)
    doc_docx.add_paragraph(
        "Except for breaches of confidentiality or gross negligence, neither party's aggregate liability "
        "under this agreement shall exceed the total fees paid by Customer during the twelve (12) months "
        "preceding the incident giving rise to liability."
    )
    doc_docx.save(str(docx_path))
    print(f"Created {docx_path}")

    # 3. Engineering Hardware & Maintenance Manual (PDF)
    eng_pdf_path = out_dir / "hardware_maintenance_manual_xr4400.pdf"
    doc_eng = fitz.open()
    page_e1 = doc_eng.new_page()
    page_e1.insert_text((50, 50), "ACME Engineering Systems — Operations Manual", fontsize=8, fontname="helv")
    page_e1.insert_text((50, 80), "Model XR-4400 Industrial Gateway Maintenance Specification", fontsize=15, fontname="hebo")
    page_e1.insert_text((50, 110), "1. Mechanical and Thermal Specifications", fontsize=12, fontname="hebo")
    pe_text = (
        "The Model XR-4400 industrial gateway requires periodic calibration. The primary mounting flange bolts "
        "must be tightened to a torque specification of exactly 45 Newton-meters (Nm) with a calibrated digital wrench. "
        "Operating ambient temperature must be maintained strictly between -20 degrees Celsius and +65 degrees Celsius."
    )
    page_e1.insert_textbox(fitz.Rect(50, 125, 550, 200), pe_text, fontsize=10, fontname="helv")
    page_e1.insert_text((50, 780), "Technical Manual XR-4400 | Page 1", fontsize=8, fontname="helv")

    doc_eng.save(str(eng_pdf_path))
    doc_eng.close()
    print(f"Created {eng_pdf_path}")


if __name__ == "__main__":
    create_demo_corpus()
